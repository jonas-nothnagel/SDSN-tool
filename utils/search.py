import glob, os, sys; sys.path.append('../udfPreprocess')

#import helper
import udfPreprocess.docPreprocessing as pre
import udfPreprocess.cleaning as clean

#import needed libraries
import seaborn as sns
from pandas import DataFrame
from sentence_transformers import SentenceTransformer, CrossEncoder, util
# from keybert import KeyBERT
from transformers import pipeline
import matplotlib.pyplot as plt
import numpy as np
import streamlit as st
import pandas as pd 
from rank_bm25 import BM25Okapi
from sklearn.feature_extraction import _stop_words
import string
from tqdm.autonotebook import tqdm
import numpy as np
import docx
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE 
import logging
logger = logging.getLogger(__name__)
import tempfile
import sqlite3
import configparser

### These are lexcial search related functions/methods#####

def bm25_tokenizer(text):
    tokenized_doc = []
    for token in text.lower().split():
        token = token.strip(string.punctuation)

        if len(token) > 0 and token not in _stop_words.ENGLISH_STOP_WORDS:
            tokenized_doc.append(token)
    return tokenized_doc
    
def bm25TokenizeDoc(paraList):
    tokenized_corpus = []
    ##########Commenting this for now########### will incorporate paragrpah splitting later.
    # for passage in tqdm(paraList):
        # if len(passage.split()) >256:
        #     # st.write("Splitting")
        #     temp  = " ".join(passage.split()[:256])
        #     tokenized_corpus.append(bm25_tokenizer(temp))
        #     temp  = " ".join(passage.split()[256:])
        #     tokenized_corpus.append(bm25_tokenizer(temp))
        # else:
        #     tokenized_corpus.append(bm25_tokenizer(passage))
    ######################################################################################33333
    for passage in tqdm(paraList):
        tokenized_corpus.append(bm25_tokenizer(passage))    

    return tokenized_corpus

def lexical_search(keyword, document_bm25):
    config = configparser.ConfigParser()
    config.read_file(open('udfPreprocess/paramconfig.cfg'))
    top_k = int(config.get('lexical_search','TOP_K'))
    bm25_scores = document_bm25.get_scores(bm25_tokenizer(keyword))
    top_n = np.argpartition(bm25_scores, -top_k)[-top_k:]
    bm25_hits = [{'corpus_id': idx, 'score': bm25_scores[idx]} for idx in top_n]
    bm25_hits = sorted(bm25_hits, key=lambda x: x['score'], reverse=True)
    return bm25_hits

@st.cache(allow_output_mutation=True)
def load_sentenceTransformer(name):
    return SentenceTransformer(name)


def semantic_search(keywordlist,paraList):
    
    ##### Sematic Search #####
    #query = "Does document contain {} issues ?".format(keyword)
    config = configparser.ConfigParser()
    config.read_file(open('udfPreprocess/paramconfig.cfg'))
    model_name = config.get('semantic_search','MODEL_NAME')

    bi_encoder = load_sentenceTransformer(model_name)
    bi_encoder.max_seq_length = int(config.get('semantic_search','MAX_SEQ_LENGTH'))     #Truncate long passages to 256 tokens
    top_k = int(config.get('semantic_search','TOP_K'))  
    document_embeddings = bi_encoder.encode(paraList, convert_to_tensor=True, show_progress_bar=False)
    question_embedding = bi_encoder.encode(keywordlist, convert_to_tensor=True)

    hits = util.semantic_search(question_embedding, document_embeddings, top_k=top_k)
    
    return hits

def show_results(keywordList):
            document = docx.Document()
            # document.add_heading('Document name:{}'.format(file_name), 2)
            section = document.sections[0]

           # Calling the footer
            footer = section.footer
        
            # Calling the paragraph already present in
        # the footer section
            footer_para = footer.paragraphs[0]
        
            font_styles = document.styles
            font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
            font_object = font_charstyle.font
            font_object.size = Pt(7)
        # Adding the centered zoned footer
            footer_para.add_run('''\tPowered by GIZ Data and the Sustainable Development Solution Network hosted at Hugging-Face spaces: https://huggingface.co/spaces/ppsingh/streamlit_dev''', style='CommentsStyle')
            document.add_heading('Your Seacrhed for {}'.format(keywordList), level=1)
            for keyword in keywordList:
          
              st.write("Results for Query: {}".format(keyword))
              para = document.add_paragraph().add_run("Results for Query: {}".format(keyword))
              para.font.size = Pt(12)
              bm25_hits, hits = search(keyword)     

              st.markdown("""
                      We will provide with 2 kind of results. The 'lexical search' and the semantic search. 
                      """)  
              # In the semantic search part we provide two kind of results one with only Retriever (Bi-Encoder) and other the ReRanker (Cross Encoder)           
              st.markdown("Top few lexical search (BM25) hits")
              document.add_paragraph("Top few lexical search (BM25) hits")

              for hit in bm25_hits[0:5]:
                  if hit['score'] > 0.00:   
                      st.write("\t Score: {:.3f}:  \t{}".format(hit['score'], paraList[hit['corpus_id']].replace("\n", " ")))
                      document.add_paragraph("\t Score: {:.3f}:  \t{}".format(hit['score'], paraList[hit['corpus_id']].replace("\n", " ")))
          
        
        
        #   st.table(bm25_hits[0:3])
          
              st.markdown("\n-------------------------\n")
              st.markdown("Top few Bi-Encoder Retrieval hits")
              document.add_paragraph("\n-------------------------\n")
              document.add_paragraph("Top few Bi-Encoder Retrieval hits")

              hits = sorted(hits, key=lambda x: x['score'], reverse=True)
              for hit in hits[0:5]:
                #  if hit['score'] > 0.45:
                  st.write("\t Score: {:.3f}:  \t{}".format(hit['score'], paraList[hit['corpus_id']].replace("\n", " ")))
                  document.add_paragraph("\t Score: {:.3f}:  \t{}".format(hit['score'], paraList[hit['corpus_id']].replace("\n", " ")))